from docx import Document
from langchain_ollama import OllamaLLM
from langchain_core.prompts import ChatPromptTemplate

class WORDDocument:
    
    def __init__(self, model_name="llama3.1"):
        self.model = OllamaLLM(model=model_name)
        self.context = ""
        self.document_text=""
        self.template = """
        answer the question below.
        Here is the word document: {document_text}
        here is the conservation history: {context}
        question: {question}
        answer:
        """
        self.prompt = ChatPromptTemplate.from_template(self.template)
        self.chain = self.prompt | self.model
    # Function to read the content of a Word document
    def read_docx(self,file_path):
        document = Document(file_path)
        full_text = []
        for para in document.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)

    # Main function


    def get_docx_response(self, question, word_file_path):
        # Ask for the path to the Word document
        file_path = word_file_path

        try:
            # Read the document content
            self.document_text = self.read_docx(file_path)

            # Get the user's question
            user_question = question
            llm= OllamaLLM(model="llama3.1")


            response = self.chain.invoke({"context": self.context, "document_text":self.document_text, "question": user_question})
            self.context += f"\nUser: {user_question}\nAI: {response}1"
            return response

        except Exception as e:
            print(f"An error occurred: {e}")

if __name__ == "__main__":
    word_model = WORDDocument()
    res= word_model.get_docx_response(word_file_path="ROLES.docx",question="what the document about?")
    print(res)

